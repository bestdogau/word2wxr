"""
Word Document Parser
Extracts structured content from .docx files including:
- Blog meta information (SEO title, description, slug, keywords)
- Headings (H1-H6)
- Paragraphs with inline formatting (bold, italic, links, superscript)
- Lists (bullet and numbered) with formatting preserved
- Tables with formatting preserved
- FAQ sections (Q:/A: pattern)
- Categories
- Hyperlinks (both Word hyperlinks and auto-detected URLs)
"""

import re
from docx import Document
from docx.oxml.ns import qn


def parse_docx(file_path):
    doc = Document(file_path)
    
    result = {
        'meta': {},
        'title': '',
        'content': [],
        'categories': [],
        'tags': [],
    }
    
    paragraphs = doc.paragraphs
    tables = doc.tables
    
    body = doc.element.body
    elements = []
    table_index = 0
    para_index = 0
    
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p' and para_index < len(paragraphs):
            elements.append(('paragraph', paragraphs[para_index]))
            para_index += 1
        elif tag == 'tbl' and table_index < len(tables):
            elements.append(('table', tables[table_index]))
            table_index += 1
    
    # Phase 1: Extract meta information
    # First, count H1s to determine document structure
    h1_positions = []
    for idx, (etype, elem) in enumerate(elements):
        if etype == 'paragraph':
            style = elem.style.name if elem.style else 'Normal'
            if style == 'Heading 1' and elem.text.strip():
                h1_positions.append(idx)
    
    i = 0
    
    if len(h1_positions) >= 2:
        # TWO+ H1s: first is meta marker, meta between them, second is title
        first_h1 = h1_positions[0]
        second_h1 = h1_positions[1]
        result['title'] = elements[second_h1][1].text.strip()
        # Parse meta between the two H1s
        for idx in range(first_h1 + 1, second_h1):
            etype, elem = elements[idx]
            if etype == 'paragraph':
                text = elem.text.strip()
                if text:
                    _parse_meta_line(text, result['meta'])
        i = second_h1 + 1
    
    elif len(h1_positions) == 1:
        # ONE H1: check for meta lines before it, H1 is the title
        h1_idx = h1_positions[0]
        result['title'] = elements[h1_idx][1].text.strip()
        # Try to parse any meta lines before the H1
        for idx in range(0, h1_idx):
            etype, elem = elements[idx]
            if etype == 'paragraph':
                text = elem.text.strip()
                if text:
                    _parse_meta_line(text, result['meta'])
        i = h1_idx + 1
    
    else:
        # NO H1s: use filename as title, start from beginning
        import os
        basename = os.path.splitext(os.path.basename(file_path))[0]
        result['title'] = basename.replace('-', ' ').replace('_', ' ').title()
        i = 0
    
    # Phase 2: Parse content blocks
    faq_mode = False
    faq_questions = []
    current_faq_q = None
    current_faq_a = None
    list_items = []
    list_is_ordered = False
    
    while i < len(elements):
        etype, elem = elements[i]
        
        if etype == 'table':
            if list_items:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            table_data = _parse_table(elem)
            result['content'].append({'type': 'table', 'headers': table_data['headers'], 'rows': table_data['rows'], 'col_count': table_data['col_count']})
            i += 1
            continue
        
        style = elem.style.name if elem.style else 'Normal'
        text = elem.text.strip()
        
        if not text:
            if list_items:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            i += 1
            continue
        
        # Categories line
        if text.startswith('Categories:'):
            if list_items:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            result['categories'] = _parse_categories_line(text)
            i += 1
            continue
        
        # FAQ section start
        if style == 'Heading 2' and text.lower() in ['faqs', 'faq', 'frequently asked questions']:
            if list_items:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            faq_mode = True
            result['content'].append({'type': 'heading', 'level': 2, 'text': text})
            i += 1
            continue
        
        # End FAQ mode on next H2
        if style == 'Heading 2' and faq_mode:
            if current_faq_q and current_faq_a:
                faq_questions.append({'question': current_faq_q, 'answer': current_faq_a})
            if faq_questions:
                result['content'].append({'type': 'rank_math_faq', 'questions': faq_questions.copy()})
                faq_questions = []
            faq_mode = False
            current_faq_q = None
            current_faq_a = None
        
        # Inside FAQ mode
        if faq_mode and style not in ['Heading 2', 'Heading 3', 'Heading 4']:
            # Format 1: Q:/A: prefixed lines
            if text.startswith('Q:') or text.startswith('Q :'):
                if current_faq_q and current_faq_a:
                    faq_questions.append({'question': current_faq_q, 'answer': current_faq_a})
                current_faq_q = text[2:].strip()
                current_faq_a = None
            elif text.startswith('A:') or text.startswith('A :'):
                current_faq_a = text[2:].strip()
            # Format 2: Alternating paragraphs — question ends with ?
            elif current_faq_q is None and text.rstrip().endswith('?'):
                # This is a question line (no Q: prefix)
                current_faq_q = text.strip()
                current_faq_a = None
            elif current_faq_q is not None and current_faq_a is None:
                # This is the answer to the previous question (regardless of ? ending)
                current_faq_a = text.strip()
            elif current_faq_q is not None and current_faq_a is not None and text.rstrip().endswith('?'):
                # Save previous Q&A pair, start new question
                faq_questions.append({'question': current_faq_q, 'answer': current_faq_a})
                current_faq_q = text.strip()
                current_faq_a = None
            elif current_faq_a is not None:
                # Continuation of answer
                current_faq_a += ' ' + text
            i += 1
            continue
        
        # Headings
        if style in ['Heading 2', 'Heading 3', 'Heading 4']:
            if list_items:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            level = int(style[-1])
            result['content'].append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue
        
        # Lists (Word 'List Paragraph' style)
        if style == 'List Paragraph':
            rich_text = _get_rich_text(elem)
            is_ordered = _is_numbered_list(elem)
            if list_items and is_ordered != list_is_ordered:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            list_is_ordered = is_ordered
            list_items.append(rich_text)
            i += 1
            continue
        
        # Bullet points using • symbol (not a Word list style, but still a list)
        if text.startswith('•') or (text.startswith('- ') and not text.startswith('---')):
            rich_text = _get_rich_text(elem)
            # Strip the bullet character
            if rich_text.startswith('•'):
                rich_text = rich_text[1:].strip()
            elif rich_text.startswith('- '):
                rich_text = rich_text[2:].strip()
            if list_items and list_is_ordered:
                result['content'].append({'type': 'list', 'style': 'ordered', 'items': list_items.copy()})
                list_items = []
            list_is_ordered = False
            list_items.append(rich_text)
            i += 1
            continue
        
        # Separator marker
        if text == '[SEPARATOR]':
            if list_items:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            result['content'].append({'type': 'separator'})
            i += 1
            continue
        
        # Button marker
        button_match = re.match(r'\[BUTTON\s+url="([^"]+)"(?:\s+style="([^"]+)")?\](.*?)\[/BUTTON\]', text)
        if button_match:
            if list_items:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            result['content'].append({'type': 'button', 'url': button_match.group(1), 'style': button_match.group(2) or 'fill', 'text': button_match.group(3)})
            i += 1
            continue
        
        # Accordion marker
        if text == '[ACCORDION]':
            if list_items:
                result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
                list_items = []
            accordion_items = []
            i += 1
            cur_t = None
            cur_c = None
            while i < len(elements):
                atype, aelem = elements[i]
                if atype == 'paragraph':
                    atext = aelem.text.strip()
                    if atext == '[/ACCORDION]':
                        if cur_t and cur_c:
                            accordion_items.append({'title': cur_t, 'content': cur_c})
                        break
                    if atext.startswith('Title:'):
                        if cur_t and cur_c:
                            accordion_items.append({'title': cur_t, 'content': cur_c})
                        cur_t = atext[6:].strip()
                        cur_c = None
                    elif atext.startswith('Content:'):
                        cur_c = atext[8:].strip()
                    elif cur_c:
                        cur_c += ' ' + atext
                i += 1
            if accordion_items:
                result['content'].append({'type': 'accordion', 'items': accordion_items})
            i += 1
            continue
        
        # Regular paragraph
        if list_items:
            result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
            list_items = []
        rich_text = _get_rich_text(elem)
        result['content'].append({'type': 'paragraph', 'text': rich_text})
        i += 1
    
    # Flush remaining
    if faq_mode and current_faq_q and current_faq_a:
        faq_questions.append({'question': current_faq_q, 'answer': current_faq_a})
    if faq_questions:
        result['content'].append({'type': 'rank_math_faq', 'questions': faq_questions})
    if list_items:
        result['content'].append({'type': 'list', 'style': 'ordered' if list_is_ordered else 'unordered', 'items': list_items.copy()})
    
    return result


def _is_numbered_list(paragraph):
    """Check if a list paragraph is numbered (ordered) or bulleted."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            numId = numPr.find(qn('w:numId'))
            if numId is not None:
                num_val = numId.get(qn('w:val'))
                try:
                    doc_part = paragraph.part
                    numbering_part = doc_part.numbering_part
                    if numbering_part is not None:
                        numbering_xml = numbering_part._element
                        # Find which abstractNum this numId references
                        for num in numbering_xml.findall(qn('w:num')):
                            if num.get(qn('w:numId')) == num_val:
                                abstractNumId_ref = num.find(qn('w:abstractNumId'))
                                if abstractNumId_ref is not None:
                                    abs_id = abstractNumId_ref.get(qn('w:val'))
                                    # Find the abstract numbering definition
                                    for abstractNum in numbering_xml.findall(qn('w:abstractNum')):
                                        if abstractNum.get(qn('w:abstractNumId')) == abs_id:
                                            # Check level 0 format
                                            for lvl in abstractNum.findall(qn('w:lvl')):
                                                if lvl.get(qn('w:ilvl')) == '0':
                                                    numFmt = lvl.find(qn('w:numFmt'))
                                                    if numFmt is not None:
                                                        fmt = numFmt.get(qn('w:val'))
                                                        if fmt in ('decimal', 'lowerLetter', 'upperLetter', 'lowerRoman', 'upperRoman'):
                                                            return True
                                                        return False
                except Exception:
                    pass
    return False


def _parse_meta_line(text, meta):
    if text.lower().startswith('title tag'):
        match = re.match(r'Title Tag.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['seo_title'] = match.group(1).strip()
    elif text.lower().startswith('meta description'):
        match = re.match(r'Meta Description.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['seo_description'] = match.group(1).strip()
    elif text.lower().startswith('url slug'):
        match = re.match(r'URL Slug.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['slug'] = match.group(1).strip().strip('/')
    elif text.lower().startswith('primary keyword'):
        match = re.match(r'Primary Keyword.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['focus_keyword'] = match.group(1).strip()
    elif text.lower().startswith('secondary keyword'):
        match = re.match(r'Secondary Keywords?.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['secondary_keywords'] = [k.strip() for k in match.group(1).split('|')]
    elif text.lower().startswith('author'):
        match = re.match(r'Author.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['author'] = match.group(1).strip()
    elif text.lower().startswith('category') and not text.lower().startswith('categories'):
        match = re.match(r'Category.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['category'] = match.group(1).strip()
    elif text.lower().startswith('tags'):
        match = re.match(r'Tags.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['tags'] = [t.strip() for t in match.group(1).split(',')]
    elif text.lower().startswith('status'):
        match = re.match(r'Status.*?:\s*(.*)', text, re.IGNORECASE)
        if match:
            meta['status'] = match.group(1).strip().lower()


def _get_rich_text(paragraph):
    """Extract rich text preserving bold, italic, hyperlinks, superscript."""
    html_parts = []
    
    for child in paragraph._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'hyperlink':
            _process_hyperlink(child, paragraph, html_parts)
        elif tag == 'r':
            _process_run(child, html_parts)
    
    result = ''.join(html_parts)
    result = re.sub(r'</strong><strong>', '', result)
    result = re.sub(r'</em><em>', '', result)
    result = re.sub(r'</strong>(\s*)<strong>', r'\1', result)
    result = re.sub(r'</em>(\s*)<em>', r'\1', result)
    result = _auto_link_urls(result)
    return result


def _process_hyperlink(hyperlink_elem, paragraph, html_parts):
    """Process a Word hyperlink element."""
    r_id = hyperlink_elem.get(qn('r:id'))
    url = ''
    if r_id and hasattr(paragraph, 'part') and paragraph.part:
        try:
            url = paragraph.part.rels[r_id].target_ref
        except (KeyError, AttributeError):
            pass
    if not url:
        anchor = hyperlink_elem.get(qn('w:anchor'))
        if anchor:
            url = f'#{anchor}'
    
    link_parts = []
    for run_elem in hyperlink_elem.findall(qn('w:r')):
        t_elem = run_elem.find(qn('w:t'))
        if t_elem is None or t_elem.text is None:
            continue
        text = t_elem.text
        rpr = run_elem.find(qn('w:rPr'))
        is_bold = False
        is_italic = False
        if rpr is not None:
            b = rpr.find(qn('w:b'))
            if b is not None:
                val = b.get(qn('w:val'))
                is_bold = val is None or val.lower() not in ('false', '0')
            it = rpr.find(qn('w:i'))
            if it is not None:
                val = it.get(qn('w:val'))
                is_italic = val is None or val.lower() not in ('false', '0')
        escaped = _escape_html(text)
        if is_bold and is_italic:
            link_parts.append(f'<strong><em>{escaped}</em></strong>')
        elif is_bold:
            link_parts.append(f'<strong>{escaped}</strong>')
        elif is_italic:
            link_parts.append(f'<em>{escaped}</em>')
        else:
            link_parts.append(escaped)
    
    link_text = ''.join(link_parts)
    if url and link_text:
        escaped_url = _escape_html(url)
        html_parts.append(f'<a href="{escaped_url}" target="_blank" rel="noreferrer noopener">{link_text}</a>')
    elif link_text:
        html_parts.append(link_text)


def _process_run(run_elem, html_parts):
    """Process a Word run element with formatting."""
    t_elem = run_elem.find(qn('w:t'))
    if t_elem is None or t_elem.text is None:
        return
    text = t_elem.text
    rpr = run_elem.find(qn('w:rPr'))
    is_bold = False
    is_italic = False
    is_superscript = False
    
    if rpr is not None:
        b = rpr.find(qn('w:b'))
        if b is not None:
            val = b.get(qn('w:val'))
            is_bold = val is None or val.lower() not in ('false', '0')
        it = rpr.find(qn('w:i'))
        if it is not None:
            val = it.get(qn('w:val'))
            is_italic = val is None or val.lower() not in ('false', '0')
        va = rpr.find(qn('w:vertAlign'))
        if va is not None and va.get(qn('w:val')) == 'superscript':
            is_superscript = True
    
    escaped = _escape_html(text)
    if is_superscript:
        escaped = f'<sup>{escaped}</sup>'
    if is_bold and is_italic:
        html_parts.append(f'<strong><em>{escaped}</em></strong>')
    elif is_bold:
        html_parts.append(f'<strong>{escaped}</strong>')
    elif is_italic:
        html_parts.append(f'<em>{escaped}</em>')
    else:
        html_parts.append(escaped)


def _auto_link_urls(html_text):
    """Auto-detect plain text URLs not already inside <a> tags."""
    if 'http' not in html_text and 'www.' not in html_text:
        return html_text
    parts = re.split(r'(<a\s[^>]*>.*?</a>)', html_text, flags=re.DOTALL)
    result = []
    for part in parts:
        if part.startswith('<a '):
            result.append(part)
        else:
            linked = re.sub(
                r'(https?://[^\s<>"\']+)',
                r'<a href="\1" target="_blank" rel="noreferrer noopener">\1</a>',
                part
            )
            result.append(linked)
    return ''.join(result)


def _parse_table(table):
    rows = []
    headers = []
    col_count = len(table.columns)
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            cell_parts = []
            for para in cell.paragraphs:
                rich = _get_rich_text(para)
                if rich.strip():
                    cell_parts.append(rich)
            cells.append(' '.join(cell_parts))
        if i == 0:
            # Strip <strong> from headers — <th> is already bold
            headers = [re.sub(r'</?strong>', '', c) for c in cells]
        else:
            rows.append(cells)
    return {'headers': headers, 'rows': rows, 'col_count': col_count}


def _parse_categories_line(text):
    categories = []
    text = re.sub(r'^Categories:\s*', '', text)
    parts = re.findall(r'([^,\(]+?)(?:\s*\(([^)]*)\))?\s*(?:,|$)', text)
    for name, url in parts:
        name = name.strip()
        if name:
            nicename = re.sub(r'[^a-z0-9]+', '-', name.lower()).strip('-')
            categories.append({'name': name, 'nicename': nicename, 'url': url.strip() if url else ''})
    return categories


def _escape_html(text):
    return (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;'))


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <file.docx>")
        sys.exit(1)
    result = parse_docx(sys.argv[1])
    print("=== META ===")
    for k, v in result['meta'].items():
        print(f"  {k}: {v}")
    print(f"\n=== TITLE ===\n  {result['title']}")
    print(f"\n=== CONTENT ({len(result['content'])} blocks) ===")
    for block in result['content']:
        btype = block['type']
        if btype == 'heading':
            print(f"  H{block['level']}: {block['text'][:80]}")
        elif btype == 'paragraph':
            print(f"  P: {block['text'][:80]}")
        elif btype == 'list':
            sl = 'OL' if block['style'] == 'ordered' else 'UL'
            print(f"  {sl} ({len(block['items'])} items): {block['items'][0][:60]}...")
        elif btype == 'table':
            print(f"  TABLE: {len(block['rows'])+1} rows x {block['col_count']} cols")
        elif btype == 'rank_math_faq':
            print(f"  FAQ: {len(block['questions'])} questions")
        elif btype == 'separator':
            print(f"  ---")
        elif btype == 'button':
            print(f"  BTN: {block['text']} -> {block['url']}")
        elif btype == 'accordion':
            print(f"  ACC: {len(block['items'])} items")
    print(f"\n=== CATEGORIES ({len(result['categories'])}) ===")
    for cat in result['categories']:
        print(f"  {cat['name']} ({cat['nicename']})")
